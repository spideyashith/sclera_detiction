from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
import os
import re

def set_academic_style(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.style.font.name = 'Times New Roman'
    paragraph.style.font.size = Pt(12)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(12)
    paragraph_format.alignment = alignment

def add_header_footer(doc):
    # Add page numbers to footer
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adding page number field
    # Reference: https://stackoverflow.com/questions/33215664/how-to-add-page-numbers-to-a-docx-file-using-python-docx
    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    add_page_number(p.add_run())

def add_heading(doc, text, level):
    h = doc.add_heading(text, level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = None # Set to black
    if level == 1:
        h.paragraph_format.space_before = Pt(36)
        h.paragraph_format.space_after = Pt(24)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(12)

def add_placeholder_space(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add a border-like space or just several empty lines
    for _ in range(12):
        doc.add_paragraph()
    
    msg = doc.add_paragraph(f"<<< PASTE {text} HERE >>>")
    msg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in msg.runs:
        run.bold = True
        run.font.size = Pt(14)
    
    for _ in range(10):
        doc.add_paragraph()

def add_paragraph_with_formatting(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, is_list=False):
    p = doc.add_paragraph()
    if is_list:
        if text.startswith('- '):
             p.style = 'List Bullet'
             text = text[2:].strip()
        else:
             p.style = 'List Number'
             text = text[text.find('. ')+2:].strip()
    
    # Process [Ref] as bold
    text = re.sub(r'(\[\d+\])', r'**\1**', text)

    # Simple regex for bold **text** and italic *text*
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)
            
    set_academic_style(p, alignment)
    return p

def add_image(doc, image_path, caption):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(5.5))
        
        cap_p = doc.add_paragraph(f"Figure: {caption}")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_academic_style(cap_p, WD_ALIGN_PARAGRAPH.CENTER)
    else:
        # If image missing, add a placeholder box
        add_placeholder_space(doc, caption.upper())

def generate_thesis():
    doc = Document()
    
    # Set Margins (1 inch = 72 points * 1)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25) # Slightly more for binding
    section.right_margin = Inches(1)

    add_header_footer(doc)

    drafts_dir = 'thesis_drafts'
    chapter_files = [
        'Front_Matter.md',
        'Chapter_1_Introduction.md',
        'Chapter_2_Literature_Review.md',
        'Chapter_3_Problem_Statement.md',
        'Chapter_4_System_Architecture.md',
        'Chapter_5_Dataset_Description.md',
        'Chapter_6_Methodology.md',
        'Chapter_7_Implementation.md',
        'Chapter_8_Experimental_Results.md',
        'Chapter_9_Dataset_Analysis.md',
        'Chapter_10_System_Interface.md',
        'Chapter_11_Limitations.md',
        'Chapter_12_Future_Work.md',
        'Chapter_13_Conclusion.md',
        'References.md',
        'Appendix_A.md'
    ]

    for filename in chapter_files:
        path = os.path.join(drafts_dir, filename)
        if not os.path.exists(path):
            continue
            
        with open(path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        in_code_block = False
        for line in lines:
            line = line.rstrip()
            if not line and not in_code_block:
                continue
            
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue

            # Handle separation of pages for front matter markers
            if line == '---':
                doc.add_page_break()
                continue

            if in_code_block:
                p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                continue

            line = line.strip()
            if not line: continue

            if line.startswith('# '):
                # Major chapters start on new page
                if filename != 'Front_Matter.md' or (line == '# CERTIFICATE' or line == '# DECLARATION' or line == '# SYNOPSIS' or line == '# ABSTRACT'):
                    doc.add_page_break()
                
                if line == '# TITLE PAGE':
                    continue 

                if line == '# CHAPTER 1: INTRODUCTION':
                     # Insert TOC and Lists
                     doc.add_page_break()
                     add_heading(doc, 'TABLE OF CONTENTS', 1)
                     p = doc.add_paragraph()
                     run = p.add_run()
                     fldChar = OxmlElement('w:fldChar')
                     fldChar.set(ns.qn('w:fldCharType'), 'begin')
                     instrText = OxmlElement('w:instrText')
                     instrText.set(ns.qn('xml:space'), 'preserve')
                     instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
                     fldChar2 = OxmlElement('w:fldChar')
                     fldChar2.set(ns.qn('w:fldCharType'), 'separate')
                     fldChar3 = OxmlElement('w:fldChar')
                     fldChar3.set(ns.qn('w:fldCharType'), 'end')
                     run._r.append(fldChar)
                     run._r.append(instrText)
                     run._r.append(fldChar2)
                     run._r.append(fldChar3)
                     
                     doc.add_page_break()
                     add_heading(doc, 'LIST OF FIGURES', 1)
                     doc.add_paragraph("Table of Figures will be generated here.")
                     doc.add_page_break()
                     add_heading(doc, 'LIST OF TABLES', 1)
                     doc.add_paragraph("List of Tables will be generated here.")
                     doc.add_page_break()
                
                add_heading(doc, line[2:], 1)
            elif line.startswith('## '):
                add_heading(doc, line[3:], 2)
            elif line.startswith('### '):
                add_heading(doc, line[4:], 3)
            elif line.startswith('[INSERT_'):
                 # Placeholder detected
                 marker = line[1:-1].replace('_HERE', '').replace('INSERT_', '').replace('_', ' ')
                 add_placeholder_space(doc, marker)
            elif line.startswith('- ') or (line[0].isdigit() and line[1:3] == '. '):
                add_paragraph_with_formatting(doc, line, is_list=True)
            elif line.startswith('|'):
                # Simple table handling
                # Since proper table parsing is complex, we just add it as a monospaced paragraph or text
                p = doc.add_paragraph(line)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            else:
                alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if filename == 'Front_Matter.md':
                     if any(marker in line for marker in ["Submitted to", "MASTER OF SCIENCE", "ASHITH JOSWA", "DR. RUBAN S", "ST ALOYSIUS", "MARCH 2026"]):
                         alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                add_paragraph_with_formatting(doc, line, alignment=alignment)
        
        # Add relevant figures per chapter
        if filename == 'Chapter_4_System_Architecture.md':
             add_image(doc, 'sclera_overlay.png', "System Visualization showing Sclera Overlay")
        elif filename == 'Chapter_8_Experimental_Results.md':
             add_image(doc, 'feature_importance.png', "Random Forest Feature Importance Analysis")
             add_image(doc, 'predicted_mask.png', "U-Net Generated Sclera Segmentation Mask")
        elif filename == 'Chapter_9_Dataset_Analysis.md':
             add_image(doc, 'gender_jaundice_comparison.png', "Jaundice Prevalence across Genders")
             add_image(doc, 'mean_b_lab_distribution.png', "Distribution of LAB 'b' content in Sclera")

    output_filename = "Academic_Thesis_Jaundice_Detection_v2.docx"
    doc.save(output_filename)
    print(f"Thesis successfully generated: {output_filename}")

if __name__ == "__main__":
    generate_thesis()
