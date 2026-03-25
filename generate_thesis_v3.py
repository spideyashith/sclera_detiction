from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
import os
import re

def set_academic_style(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.style.font.name = 'Times New Roman'
    paragraph.style.font.size = Pt(12)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(12)
    paragraph_format.alignment = alignment

def add_header_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def create_element(name): return OxmlElement(name)
    def create_attribute(element, name, value): element.set(ns.qn(name), value)
    
    def add_page_number(run):
        run._r.append(create_element('w:fldChar'))
        run._r[-1].set(ns.qn('w:fldCharType'), 'begin')
        instr = create_element('w:instrText')
        instr.set(ns.qn('xml:space'), 'preserve')
        instr.text = "PAGE"
        run._r.append(instr)
        run._r.append(create_element('w:fldChar'))
        run._r[-1].set(ns.qn('w:fldCharType'), 'end')

    add_page_number(p.add_run())

def add_heading(doc, text, level):
    h = doc.add_heading(text, level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = None
    if level == 1:
        h.paragraph_format.space_before = Pt(36)
        h.paragraph_format.space_after = Pt(24)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(12)

def add_placeholder_space(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(10): doc.add_paragraph()
    msg = doc.add_paragraph(f"<<< {text} >>>")
    msg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in msg.runs:
        run.bold = True
        run.font.size = Pt(14)
    for _ in range(10): doc.add_paragraph()

def add_paragraph_with_formatting(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, is_list=False):
    p = doc.add_paragraph()
    if is_list:
        if text.startswith('- '):
             p.style = 'List Bullet'
             text = text[2:].strip()
        else:
             p.style = 'List Number'
             text = text[text.find('. ')+2:].strip()
    
    text = re.sub(r'(\[\d+\])', r'**\1**', text)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)
    set_academic_style(p, alignment)

def add_image(doc, image_path, caption):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(5.0))
        cap_p = doc.add_paragraph(f"Figure: {caption}")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_academic_style(cap_p, WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_placeholder_space(doc, f"INSERT {caption.upper()} HERE")

def generate_thesis():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)
    add_header_footer(doc)

    drafts_dir = 'thesis_v3'
    chapter_files = [
        'Front_Matter.md',
        'Chapter_1_Introduction.md',
        'Chapter_2_Literature_Review.md',
        'Chapter_3_Methodology.md',
        'Chapter_4_Results.md',
        'Chapter_5_Conclusion.md',
        'References.md',
        'Appendix_A.md'
    ]

    for filename in chapter_files:
        path = os.path.join(drafts_dir, filename)
        if not os.path.exists(path): continue
        with open(path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        in_code_block = False
        for line in lines:
            line = line.rstrip()
            if not line and not in_code_block: continue
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue
            if line == '---':
                doc.add_page_break()
                continue
            if in_code_block:
                p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                continue
            line = line.strip()
            if not line: continue
            if line.startswith('# '):
                if filename != 'Front_Matter.md' or any(m in line for m in ['CERTIFICATE', 'DECLARATION', 'SYNOPSIS', 'ABSTRACT']):
                    doc.add_page_break()
                if line == '# CHAPTER 1': # Insert TOC, LOF, LOT before Ch 1
                    doc.add_page_break()
                    add_heading(doc, 'TABLE OF CONTENTS', 1)
                    p = doc.add_paragraph()
                    run = p.add_run()
                    fldChar = OxmlElement('w:fldChar')
                    fldChar.set(ns.qn('w:fldCharType'), 'begin')
                    instrText = OxmlElement('w:instrText'); instrText.set(ns.qn('xml:space'), 'preserve')
                    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
                    run._r.append(fldChar); run._r.append(instrText)
                    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(ns.qn('w:fldCharType'), 'separate')
                    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(ns.qn('w:fldCharType'), 'end')
                    run._r.append(fldChar2); run._r.append(fldChar3)
                    
                    doc.add_page_break()
                    add_heading(doc, 'LIST OF FIGURES', 1)
                    doc.add_paragraph("List of Figures (to be updated in Word)")
                    doc.add_page_break()
                    add_heading(doc, 'LIST OF TABLES', 1)
                    doc.add_paragraph("List of Tables (to be updated in Word)")
                    doc.add_page_break()
                add_heading(doc, line[2:], 1)
            elif line.startswith('## '): add_heading(doc, line[3:], 2)
            elif line.startswith('### '): add_heading(doc, line[4:], 3)
            elif line.startswith('[INSERT_'):
                marker = line[1:-1].replace('_HERE', '').replace('INSERT_', '').replace('_', ' ')
                add_placeholder_space(doc, marker)
            elif line.startswith('- ') or (line[0].isdigit() and line[1:3] == '. '):
                add_paragraph_with_formatting(doc, line, is_list=True)
            elif line.startswith('|'):
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            else:
                alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if filename == 'Front_Matter.md' and any(m in line for m in ["Submitted to", "MASTER OF SCIENCE", "ASHITH JOSWA", "DR. RUBAN S", "ST ALOYSIUS", "MARCH 2026"]):
                    alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_paragraph_with_formatting(doc, line, alignment=alignment)
        
        # Add Figures
        if filename == 'Chapter_3_Methodology.md':
            add_image(doc, 'sclera_overlay.png', "System Visualization showing Sclera Overlay")
        elif filename == 'Chapter_4_Results.md':
            add_image(doc, 'feature_importance.png', "Random Forest Feature Importance Analysis")
            add_image(doc, 'predicted_mask.png', "U-Net Generated Sclera Segmentation Mask")
            add_image(doc, 'mean_b_lab_distribution.png', "Distribution of LAB 'b+' content in Sclera")

    output_filename = "Sclera_Detection_Final_Documentation_90Pages.docx"
    doc.save(output_filename)
    print(f"Thesis successfully generated: {output_filename}")

if __name__ == "__main__":
    generate_thesis()
